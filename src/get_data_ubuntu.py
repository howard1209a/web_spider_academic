from selenium import webdriver
from selenium.webdriver.chrome.service import Service
import tempfile
from bs4 import BeautifulSoup
import shutil
from translate import baidu_api
import re
from docx import Document


class Spider:
    # 配置WebDriver，使用Chrome浏览器
    def __init__(self):
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')  # 不弹出浏览器界面
        options.add_argument('--disable-gpu')  # 禁用GPU加速
        options.add_argument("--no-sandbox")  # 禁用沙箱

        # 创建一个临时目录来存放 Chrome 用户数据
        self.user_data_dir = tempfile.mkdtemp()
        options.add_argument(f"--user-data-dir={self.user_data_dir}")

        # 下述为docker镜像中的chromedriver路径
        chromedriver_path = '/usr/local/bin/chromedriver-linux64/chromedriver'
        service = Service(executable_path=chromedriver_path)

        # 启动浏览器
        self.driver = webdriver.Chrome(service=service, options=options)

        # 下述为ubuntu 24.04LTS Desktop中火狐浏览器真实header
        headers = {
            "Host": "www.sciencedirect.com",
            "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:125.0) Gecko/20100101 Firefox/125.0",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Cookie": "s_sess=%20s_ppvl%3D%3B%20s_ppv%3Dsd%25253Abrowse%25253Ajournal%25253Aissue%252C8%252C8%252C433%252C1214%252C433%252C1280%252C800%252C1%252CP%3B",
            "Upgrade-Insecure-Requests": "1",
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "none",
            "Sec-Fetch-User": "?1"
        }

        # 使用 Chrome DevTools Protocol 设置请求头
        for header, value in headers.items():
            self.driver.execute_cdp_cmd('Network.setUserAgentOverride', {
                'userAgent': headers['User-Agent'],
                'acceptLanguage': headers['Accept-Language'],
                'platform': 'Windows'
            })
            self.driver.execute_cdp_cmd('Network.setExtraHTTPHeaders', {
                'headers': {header: value for header, value in headers.items() if header != 'User-Agent'}
            })

    def get_url(self, url):
        # 浏览器发起get请求
        self.driver.get(url)

        # 等待页面加载完成
        self.driver.implicitly_wait(10)

        # BeautifulSoup 解析page_source响应体数据
        return BeautifulSoup(self.driver.page_source, 'html.parser')

    # https://onlinelibrary.wiley.com/action/doSearch?SeriesKey=14678667&sortBy=Earliest
    def get_wiley_single_page(self, latest_document, page, document_writer):
        soup = self.get_url(
            "https://onlinelibrary.wiley.com/action/doSearch?SeriesKey=14678667&sortBy=Earliest&startPage=" + str(
                page))

        # 获取所有的 <div class="item__body"> 标签，每个标签代表一个文献
        items = soup.find_all('div', class_='item__body')

        should_end = False

        # 遍历每个文献，提取相关信息
        for item in items:
            # 1. 获取文献题目
            title = item.find('a', class_='publication_title visitable').text.strip()

            # 2. 获取 DOI
            doi = item.find('input', {'type': 'hidden'}).get('value')

            # 将 DOI 转换为标准的链接格式
            doi_url = f"https://doi.org/{doi}"

            if title != latest_document:
                # 写入文献信息到 Word 文档
                document_writer.add_paragraph(f"文献题目: {title}")
                document_writer.add_paragraph(f"文献中文题目: {baidu_api(title)}")
                document_writer.add_paragraph(f"DOI: {doi_url}")
                document_writer.add_paragraph("-" * 50)
            else:
                should_end = True
                break

        return should_end

    # https://www.sciencedirect.com/journal/automation-in-construction/vol/171/suppl/C
    def get_sciencedirect_single_volume(self, journal, volume, latest_document, document_writer):
        page_count = self.get_sciencedirect_page_count(journal, volume)
        for i in range(page_count, 0, -1):
            if self.get_sciencedirect_single_page(journal, volume, i, latest_document, document_writer):
                return True
        return False

    def get_sciencedirect_page_count(self, journal, volume):
        soup = self.get_url(
            "https://www.sciencedirect.com/journal/" + journal + "/vol/" + str(volume) + "/suppl/C?page=1")

        page_info_str_item = soup.find('span',
                                       class_='pagination-pages-label text-xs u-margin-s-left-from-sm u-margin-s-right-from-sm')
        if page_info_str_item is None:
            return 1

        page_info_str = page_info_str_item.text
        # 使用正则表达式提取所有数字
        match = re.search(r'(\d+)$', page_info_str)
        if match:
            # 获取匹配的数字
            print("+++++++++++++++++++++++++++++++++++" + str(
                int(match.group(1))) + "+++++++++++++++++++++++++++++++++++")
            return int(match.group(1))
        else:
            raise Exception("page count parse fail")

    # https://www.sciencedirect.com/journal/automation-in-construction/vol/171/suppl/C
    def get_sciencedirect_single_page(self, journal, volume, page, latest_document, document_writer):
        soup = self.get_url(
            "https://www.sciencedirect.com/journal/" + journal + "/vol/" + str(volume) + "/suppl/C?page=" + str(page))

        items = soup.find_all('li', class_='js-article-list-item article-item u-padding-xs-top u-margin-l-bottom')

        should_end = False

        # 遍历每个文献，提取相关信息
        for item in reversed(items):
            # 1. 获取文献题目
            title = item.find('span', class_='js-article-title text-l').text.strip()

            # 2. 获取 DOI
            doi = item.find('div', hidden=True).get('value')

            if title != latest_document:
                document_writer.add_paragraph(f"文献题目: {title}")
                document_writer.add_paragraph(f"文献中文题目: {baidu_api(title)}")
                document_writer.add_paragraph(f"DOI: {doi}")
                document_writer.add_paragraph("-" * 50)
            else:
                should_end = True
                break

        return should_end

    def close(self):
        # 关闭浏览器
        self.driver.quit()

        # 删除临时目录及其内容
        shutil.rmtree(self.user_data_dir)
